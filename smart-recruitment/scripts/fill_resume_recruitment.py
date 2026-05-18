# -*- coding: utf-8 -*-
"""Fill 智能招聘系统 section in resume docx to match 书海平台 format."""
from pathlib import Path

from docx import Document


def main() -> None:
    src = Path(r"C:\Users\86134\Desktop\不常用\简历\简历\微派-全栈.docx")
    # ASCII filename avoids Windows console / zip mojibake when cwd encoding varies.
    dst = Path(__file__).resolve().parent.parent / "weipai-fullstack-recruitment-filled.docx"

    d = Document(str(src))

    d.paragraphs[18].text = "智能招聘系统"

    tech = d.paragraphs[19]
    tech.text = (
        "技术栈: Go（Gin）+ gRPC + Protocol Buffers + JWT + GORM + MySQL + "
        "AWS SDK for Go v2（S3 兼容 OSS 预签名）+ CloudWeGo Eino（通义千问）+ bcrypt + "
        "React 19 + TypeScript + Vite"
    )
    tech.style = d.styles["Normal"]

    intro = d.paragraphs[20]
    intro.text = (
        "项目简介:面向企业与候选人的双端智能招聘平台。HTTP 网关对外提供 REST，业务集中在 gRPC 服务；"
        "支持 HR 发布与管理岗位、候选人维护结构化档案与在线投递，简历经浏览器直传私有对象存储；"
        "HR 端集成「先查库、再对话」的招聘助手，提升岗位与投递数据的查阅与问答效率。"
    )
    intro.style = d.styles["Normal"]

    core_header = d.paragraphs[21]
    blank = core_header.insert_paragraph_before("")
    blank.style = d.styles["Body Text"]

    bullets = [
        "·分层与契约: 采用 Gin HTTP 网关与 gRPC 业务服务拆分，Protobuf 定义 RPC 契约；Web 层仅做鉴权与转发，业务 SQL 与权限校验集中在 Logic，边界清晰、便于协作与迭代。",
        "·登录认证: 使用 JWT 承载用户 ID 与角色；网关解析后将身份信息写入 gRPC Metadata 透传至 Logic，统一完成 RPC 鉴权与资源归属校验（如岗位仅创建者可操作），实现无状态入口与可信调用链。",
        "·私有存储与简历安全: 基于 S3 兼容接口生成 PUT/GET 预签名 URL，浏览器直传私有 OSS；针对阿里云 OSS 适配 virtual-hosted 预签名，规避 path-style 403；Confirm 阶段读取对象头字节校验 PDF/DOC/DOCX 魔数，配合后缀与 Content-Type 白名单降低恶意文件风险。",
        "·HR 招聘助手: 先由 MySQL 聚合投递与岗位等业务事实，再通过 CloudWeGo Eino 调用通义千问生成自然语言回答；未配置模型 Key 时降级为数据库汇总文本；对话成对持久化并支持历史分页加载。",
        "·数据与投递规则: 表结构由 schema.sql 管理，运行时禁用 GORM AutoMigrate；投递结合档案完整性校验与库表唯一约束，在服务端强校验防止重复投递。",
    ]
    for i, txt in enumerate(bullets):
        p = d.paragraphs[23 + i]
        p.text = txt
        p.style = d.styles["Normal"] if i == 0 else d.styles["Body Text"]

    dst.parent.mkdir(parents=True, exist_ok=True)
    try:
        d.save(str(src))
        print("saved:", src)
    except OSError:
        d.save(str(dst))
        print("original locked or not writable; saved copy:", dst)


if __name__ == "__main__":
    main()
